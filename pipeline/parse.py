"""문서 파서 — docx/pdf/md/xlsx/이미지를 텍스트/구조로 변환.

이미지는 운영에서 HCX Vision OCR 로 처리(여기서는 자리표시). 나머지는 로컬 처리.
doc_type 은 파일명 규칙으로 추론한다.
"""
from __future__ import annotations

from pathlib import Path

SOURCE_DIR = Path(__file__).resolve().parent.parent / "souce"


def infer_doc_type(name: str) -> str:
    if "FAQ" in name:
        return "FAQ"
    if "운영기준" in name or "운영조건" in name or "상품운영" in name:
        return "운영기준"
    if "심사" in name or "규정" in name:
        return "규정"
    return "기타"


def infer_product(name: str) -> str:
    for key in ("중고리스", "중고승용", "재고금융", "중형트럭", "운영자금"):
        if key in name:
            return key
    return "공통"


def parse_xlsx(path: Path) -> list[dict]:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    header = [str(c).strip() if c else "" for c in rows[0]]
    out = []
    for r in rows[1:]:
        rec = {header[i]: (str(r[i]).replace("<br>", "\n").strip() if r[i] is not None else "")
               for i in range(min(len(header), len(r)))}
        if any(rec.values()):
            out.append(rec)
    return out


def parse_pdf(path: Path) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_md(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def parse_file(path: Path) -> dict:
    suffix = path.suffix.lower()
    meta = {"source_file": path.name, "doc_type": infer_doc_type(path.name),
            "product": infer_product(path.name)}
    if suffix == ".xlsx":
        return {"meta": meta, "kind": "rows", "rows": parse_xlsx(path)}
    if suffix == ".pdf":
        return {"meta": meta, "kind": "text", "text": parse_pdf(path)}
    if suffix == ".docx":
        return {"meta": meta, "kind": "text", "text": parse_docx(path)}
    if suffix == ".md":
        return {"meta": meta, "kind": "text", "text": parse_md(path)}
    if suffix in (".png", ".jpg", ".jpeg"):
        return {"meta": meta, "kind": "image", "text": "[이미지 — 운영 시 HCX Vision OCR 처리]"}
    return {"meta": meta, "kind": "text", "text": ""}


def parse_all(source_dir: Path = SOURCE_DIR) -> list[dict]:
    docs = []
    for p in sorted(source_dir.iterdir()):
        if p.is_file() and p.suffix.lower() in (".xlsx", ".pdf", ".docx", ".md"):
            docs.append(parse_file(p))
    return docs
